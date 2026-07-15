import streamlit as st
import docx
import qrcode
from io import BytesIO

# 1. إعدادات الصفحة بواجهة عصرية ونظيفة
st.set_page_config(
    page_title="مكتب المحاماة الذكي",
    page_icon="⚖️",
    layout="centered"
)

# تصميم هيدر احترافي بـ CSS منسق
st.markdown("""
    <style>
    .main-title {
        text-align: right; 
        color: #1E3A8A; 
        font-family: 'Cairo', sans-serif;
        margin-bottom: 5px;
    }
    .sub-title {
        text-align: right; 
        color: #4B5563; 
        font-size: 16px;
        margin-bottom: 25px;
    }
    .status-box {
        text-align: right;
        direction: rtl;
    }
    </style>
""", unsafe_allow_html=True)

st.markdown("<h1 class='main-title'>⚖️ نظام توليد الـ QR كود الذكي</h1>", unsafe_allow_html=True)
st.markdown("<p class='sub-title'>ارفع ملف القضية بصيغة Word (.docx) واضغط على زر التوليد لصنع الـ QR كود فوراً.</p>", unsafe_allow_html=True)
st.markdown("---")

# دالة استخراج النصوص من الفقرات والجداول بدقة
def extract_clean_text(file):
    doc = docx.Document(file)
    extracted_elements = []
    
    for paragraph in doc.paragraphs:
        cleaned_text = paragraph.text.strip()
        if cleaned_text:
            extracted_elements.append(cleaned_text)
            
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cleaned_cell = cell.text.strip()
                if cleaned_cell and cleaned_cell not in extracted_elements:
                    extracted_elements.append(cleaned_cell)
                    
    return "\n".join(extracted_elements)

# 2. واجهة رفع الملف (مبسطة وأنيقة)
uploaded_file = st.file_uploader("اختر ملف الوورد الخاص بالقضية", type=["docx"])

if uploaded_file is not None:
    # زرار التوليد الصريح بـ ستايل واضح
    st.markdown("<br>", unsafe_allow_html=True)
    generate_btn = st.button("⚡ توليد الـ QR كود الآن", use_container_width=True)
    
    if generate_btn:
        try:
            with st.spinner("جاري قراءة الملف وتجهيز الكود..."):
                raw_text = extract_clean_text(uploaded_file)
                
            if not raw_text.strip():
                st.error("⚠️ الملف المرفوع فارغ أو لا يحتوي على نصوص قابلة للقراءة.")
            else:
                # قص النص لـ 1000 حرف لضمان ثبات الـ QR وعدم خروج أخطاء حجم نهائياً
                max_limit = 1000
                is_truncated = False
                
                if len(raw_text) > max_limit:
                    qr_content = raw_text[:max_limit] + "\n... [تم اختصار باقي النص]"
                    is_truncated = True
                else:
                    qr_content = raw_text

                # توليد الـ QR Code
                qr_generator = qrcode.QRCode(
                    version=None,
                    error_correction=qrcode.constants.ERROR_CORRECT_L,
                    box_size=10,
                    border=4
                )
                qr_generator.add_data(qr_content)
                qr_generator.make(fit=True)
                
                qr_image = qr_generator.make_image(fill_color="#1E3A8A", back_color="white") # لون أزرق غامق كحلي فخم للـ QR بدل الأسود التقليدي
                
                image_buffer = BytesIO()
                qr_image.save(image_buffer, format="PNG")
                binary_image = image_buffer.getvalue()

                # عرض النتيجة بتنسيق نظيف جداً وسط الصفحة
                st.markdown("---")
                st.success("✅ تم توليد الـ QR كود بنجاح!")
                
                col_left, col_center, col_right = st.columns([1, 2, 1])
                with col_center:
                    st.image(binary_image, caption="اسحب الصورة أو امسحها بموبايلك", use_container_width=True)
                    
                    # زر التحميل المباشر تحت الصورة
                    st.download_button(
                        label="💾 تحميل صورة الـ QR بجودة عالية",
                        data=binary_image,
                        file_name=f"QR_{uploaded_file.name.replace('.docx', '')}.png",
                        mime="image/png",
                        use_container_width=True
                    )
                    
                # عرض اختياري للنص المستخرج أسفل الصفحة لو حابب تراجعه
                with st.expander("👁️ معاينة النص الذي تم تشفيره داخل الـ QR"):
                    st.text(raw_text)
                    if is_truncated:
                        st.warning(f"تنبيه: تم أخذ أول {max_limit} حرف فقط للحفاظ على جودة وسرعة قراءة الـ QR بالكاميرا.")
                    
        except Exception as error:
            st.error(f"❌ حدث خطأ غير متوقع: {str(error)}")

        if not text:
            st.warning("الملف لا يحتوي على نص.")
            st.stop()

        limit = 1800
        qr_text = text[:limit]
        if len(text) > limit:
            st.info("تم اختصار النص لضمان سهولة قراءة QR.")

        with st.expander("معاينة النص"):
            st.text(text)

        qr = qrcode.QRCode(
            version=None,
            error_correction=qrcode.constants.ERROR_CORRECT_H,
            box_size=14,
            border=6,
        )
        qr.add_data(qr_text)
        qr.make(fit=True)
        img = qr.make_image(fill_color="black", back_color="white")

        buf = BytesIO()
        img.save(buf, format="PNG")
        data = buf.getvalue()

        st.success("تم إنشاء QR بنجاح.")
        st.image(data, use_container_width=True)
        st.download_button(
            "تحميل QR",
            data=data,
            file_name=uploaded.name.replace(".docx","")+"_QR.png",
            mime="image/png",
            use_container_width=True,
        )
    except PackageNotFoundError:
        st.error("الملف ليس مستند DOCX صالح.")
    except PermissionError:
        st.error("الملف محمي أو لا يمكن قراءته.")
    except Exception as e:
        st.error(f"حدث خطأ: {e}")
